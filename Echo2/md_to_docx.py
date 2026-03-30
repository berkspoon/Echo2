"""Convert markdown files to styled Word documents."""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def parse_md_to_docx(md_path: str, docx_path: str):
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Heading styles
    for i in range(1, 5):
        h = doc.styles[f"Heading {i}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            level = min(level, 4)
            text = line.lstrip("#").strip()
            p = doc.add_heading(text, level=level)
            i += 1
            continue

        # Table detection
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
            table_lines = []
            while i < len(lines) and "|" in lines[i].strip():
                table_lines.append(lines[i].strip())
                i += 1

            # Parse header
            headers = [c.strip() for c in table_lines[0].split("|") if c.strip()]
            # Skip separator line (index 1)
            rows = []
            for tl in table_lines[2:]:
                cells = [c.strip() for c in tl.split("|") if c.strip()]
                # Pad if needed
                while len(cells) < len(headers):
                    cells.append("")
                rows.append(cells[:len(headers)])

            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Header row
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)

            # Data rows
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = val
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

            doc.add_paragraph()  # spacing after table
            continue

        # Checkbox list items
        if re.match(r"^- \[[ x]\] ", line):
            text = re.sub(r"^- \[[ x]\] ", "", line)
            checked = line.startswith("- [x]")
            prefix = "\u2611 " if checked else "\u2610 "
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, prefix + text)
            i += 1
            continue

        # Bullet list
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, text)
            i += 1
            continue

        # Code block
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```

            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            # Light gray background via shading
            shading = p.paragraph_format
            p_fmt = p._element
            pPr = p_fmt.get_or_add_pPr()
            shd = pPr.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "F0F0F0"
            })
            pPr.append(shd)
            continue

        # Bold metadata lines like **Key:** Value
        if line.startswith("**"):
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            i += 1
            continue

        # Regular paragraph (may span multiple lines)
        para_text = line
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith("-") and not lines[i].startswith("*") and not lines[i].startswith("|") and not lines[i].startswith("```") and not re.match(r"^\d+\.\s", lines[i]):
            para_text += " " + lines[i].strip()
            i += 1

        p = doc.add_paragraph()
        add_formatted_text(p, para_text)

    doc.save(docx_path)
    print(f"  Created: {docx_path}")


def add_formatted_text(paragraph, text):
    """Parse inline markdown (bold, italic, code) and add formatted runs."""
    # Pattern: **bold**, *italic*, `code`, [text](url) — just render text for links
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\(.+?\))"
    parts = re.split(pattern, text)

    i = 0
    while i < len(parts):
        part = parts[i]
        if part is None:
            i += 1
            continue

        # Check if this is a full match group
        if i + 1 < len(parts) and parts[i + 1] is not None and re.match(r"^\*\*(.+)\*\*$", part):
            run = paragraph.add_run(parts[i + 1])
            run.bold = True
            run.font.name = "Calibri"
            i += 5  # skip all capture groups
        elif i + 2 < len(parts) and parts[i + 2] is not None and re.match(r"^\*(.+)\*$", part):
            run = paragraph.add_run(parts[i + 2])
            run.italic = True
            run.font.name = "Calibri"
            i += 5
        elif i + 3 < len(parts) and parts[i + 3] is not None and re.match(r"^`(.+)`$", part):
            run = paragraph.add_run(parts[i + 3])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            i += 5
        elif i + 4 < len(parts) and parts[i + 4] is not None and re.match(r"^\[.+\]\(.+\)$", part):
            run = paragraph.add_run(parts[i + 4])
            run.font.name = "Calibri"
            i += 5
        else:
            if part:
                run = paragraph.add_run(part)
                run.font.name = "Calibri"
            i += 1


if __name__ == "__main__":
    base = Path("y:/VibeCode/Echo2/Interval Fund Data Extractor")
    files = [
        ("interval-fund-extractor-executive-summary.md", "Interval Fund Data Extractor - Executive Summary.docx"),
        ("interval-fund-extractor-PRD.md", "Interval Fund Data Extractor - PRD.docx"),
        ("interval-fund-extractor-prompts.md", "Interval Fund Data Extractor - Build Prompts.docx"),
    ]

    out_dir = base
    out_dir.mkdir(exist_ok=True)

    print("Converting markdown to Word...\n")
    for md_name, docx_name in files:
        md_path = base / md_name
        docx_path = out_dir / docx_name
        if md_path.exists():
            parse_md_to_docx(str(md_path), str(docx_path))
        else:
            print(f"  Skipped (not found): {md_name}")

    print("\nDone! Files are in:", out_dir)
